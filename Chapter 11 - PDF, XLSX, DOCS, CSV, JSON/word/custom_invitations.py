# custom_invitations.py - generates invitations with provided names


import docx, os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

initial_doc = docx.Document()

guests = open('guests.txt', 'r')

count = 0

for i in guests:
    if i:
        print(type(i))
        print(i[:-1])
        initial_doc.add_paragraph('It would be a pleasure to have the company of')
        initial_doc.add_paragraph(i[:-1])
        par = initial_doc.add_paragraph('at')
        par.add_run(' 11010 Memory Lane on the Evening of')
        initial_doc.add_paragraph('April 1st')
        par2 = initial_doc.add_paragraph('at')
        par2.add_run(" 7 o'clock'")

        initial_doc.paragraphs[count].runs[0].italic = True
        initial_doc.paragraphs[count].runs[0].font.name = 'Calibri'
        initial_doc.paragraphs[count].runs[0].font.size = Pt(22)
        count += 1
        initial_doc.paragraphs[count].runs[0].bold = True
        initial_doc.paragraphs[count].runs[0].font.name = 'Times New Roman'
        initial_doc.paragraphs[count].runs[0].font.size = Pt(24)
        count += 1
        initial_doc.paragraphs[count].runs[0].italic = True
        initial_doc.paragraphs[count].runs[0].font.name = 'Calibri'
        initial_doc.paragraphs[count].runs[0].font.size = Pt(22)
        initial_doc.paragraphs[count].runs[0].underline = True
        initial_doc.paragraphs[count].runs[1].italic = True
        initial_doc.paragraphs[count].runs[1].font.name = 'Calibri'
        initial_doc.paragraphs[count].runs[1].font.size = Pt(22)
        count += 1
        initial_doc.paragraphs[count].style = 'Normal'
        initial_doc.paragraphs[count].runs[0].font.name = 'Times New Roman'
        initial_doc.paragraphs[count].runs[0].font.size = Pt(22)
        count += 1
        initial_doc.paragraphs[count].runs[0].italic = True
        initial_doc.paragraphs[count].runs[0].font.name = 'Calibri'
        initial_doc.paragraphs[count].runs[0].font.size = Pt(22)
        initial_doc.paragraphs[count].runs[0].underline = True
        initial_doc.paragraphs[count].runs[1].italic = True
        initial_doc.paragraphs[count].runs[1].font.name = 'Calibri'
        initial_doc.paragraphs[count].runs[1].font.size = Pt(22)
        initial_doc.paragraphs[count].runs[1].add_break(docx.enum.text.WD_BREAK.PAGE)
        count += 1

for i in range(0, len(initial_doc.paragraphs)):
    initial_doc.paragraphs[i].alignment = WD_ALIGN_PARAGRAPH.CENTER
initial_doc.save('test.docx')


